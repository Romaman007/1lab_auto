from bs4 import BeautifulSoup
from docx import Document
import requests as req
import re
import pdfplumber
import aspose.words as aw

class Parser:
    def __init__(self,name):
        super().__init__()
        if name[-3:]=='pdf':
            self.pdf = pdfplumber.open(name)
        elif name[-3:]=='ocx':
            self.docx  = Document(name)
        elif name[-3:]=='doc':
            doc = aw.Document(name)
            doc.save('1_1.docx')
            self.docx=Document(r'1_1.docx')

        else:
            self.rep=req.get(name)
            self.soup=BeautifulSoup(self.rep.text,'html.parser')
            self.head_url = re.search('(https...[^/]*|http...[^/]*)',name)[0]
        self.links = []
        self.text=''
        self.links_img =[]
        self.tables =[]



    def get_text_html(self,form=True):
        self.text = self.soup.get_text()
        if form:
            self.text = re.sub('\n+','\n',self.text)
        else:
            self.text = re.sub('\s+',' ',self.text)
        return self.text


    def get_links_html(self):
        for link in self.soup.findAll(['a', 'link']):
            link = link.extract().get('href')
            try:
                if (link[0]=='h'):
                    self.links.append(link)
                elif (link[0]=='/'):
                    self.links.append(self.head_url+link)
            except:
                pass
        return self.links


    def get_img_html(self):
        for img in self.soup.findAll(['a','link']):
            img = img.extract().get('href')
            try:
                if (img[-3:]=='png')|(img[-3:]=='jpg')|(img[-3:]=='svg')|(img[-3:]=='peg')|(img[-3:]=='ico')|(img[-3:]=='bmp')|(img[-3:]=='gif')|(img[-3:]=='jpg'):
                    self.links_img.append(img)
            except:
                pass
        for img in self.soup.findAll(['img']):
            img = img.extract().get('src')
            self.links_img.append(self.head_url+img)

        for i in range(len(self.links_img)):
            if (self.links_img[i][0]=='/')&(self.links_img[i][1]!='/'):
                self.links_img[i] = self.head_url+self.links_img[i]
            elif (self.links_img[i][0] != '/') & (self.links_img[i][0:4] != 'http'):
                self.links_img[i] = self.head_url + '/' + self.links_img[i]
        return self.links_img


    def get_text_pdf(self):
        self.text=''
        for page in self.pdf.pages:
            self.text +='\n'+page.extract_text()
        return self.text


    def get_links_pdf(self):
        self.links = re.findall("http:\/\/?[\w-]{1,32}\.[\w-]{1,32}[^\s@,]*|https:\/\/?[\w-]{1,32}\.[\w-]{1,32}[^\s@,]*",self.text)
        return self.links


    def get_tables_pdf(self,page_with_table):
        self.tables.append(self.pdf.pages[page_with_table].extract_table())
        return self.tables


    def get_text_docx(self):
        self.text=''
        for par in self.docx.paragraphs:
            self.text+='\n'+par.text
        return self.text


    def get_links_docx(self):
        if self.text =='':
            self.get_text_docx()
        self.links = re.findall("http:\/\/?[\w-]{1,32}\.[\w-]{1,32}[^\s@,]*|https:\/\/?[\w-]{1,32}\.[\w-]{1,32}[^\s@,]*",self.text)
        return self.links

    def get_links_doc(self):
        if self.text =='':
            self.get_text_doc()
        self.links = re.findall("http:\/\/?[\w-]{1,32}\.[\w-]{1,32}[^\s@,]*|https:\/\/?[\w-]{1,32}\.[\w-]{1,32}[^\s@,]*",self.text)
        return self.links


    def get_tables_docx_doc(self):
        self.tables=[]
        for table in self.docx.tables:
            tablel=[]
            for row in table.rows:
                roww=[]
                for cell in row.cells:
                    roww.append(cell.text)
                tablel.append(roww)
            self.tables.append(tablel)
        return self.tables

    def get_text_doc(self):
        self.text=''
        flag=1
        for par in self.docx.paragraphs:
            if flag==1:
                flag=0
            else:
                self.text+='\n'+par.text
        return self.text


